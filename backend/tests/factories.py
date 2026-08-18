"""Byte-level fixtures for the extraction tests.

Building real files rather than mocking the parsers: the whole point of the
extraction stage is that it survives contact with actual PDF and OOXML
containers, and a mocked pypdf would prove nothing about that.
"""
import io

from docx import Document as DocxDocument


#: comfortably above extraction.MIN_PDF_TEXT, so a single default page is a
#: real document rather than something the OCR guard rejects
_PDF_TEXT = ("Refund policy: customers may request refunds within thirty days "
             "of the original purchase date.")


def pdf_bytes(text: str | None = _PDF_TEXT, pages: int = 1) -> bytes:
    """A minimal but genuinely valid PDF.

    text=None produces a page with no text operators at all — which is what a
    scan looks like to pypdf: it parses fine and returns "".
    """
    contents, fonts = [], b"<< /Font << /F1 %d 0 R >> >>"
    for _ in range(pages):
        stream = (b"BT /F1 12 Tf 72 720 Td (" + text.encode() + b") Tj ET"
                  if text else b"")
        contents.append(stream)

    n_pages = len(contents)
    font_obj = 3 + n_pages * 2
    kids = b" ".join(b"%d 0 R" % (3 + i * 2) for i in range(n_pages))

    objs: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [" + kids + b"] /Count %d >>" % n_pages,
    ]
    for i, stream in enumerate(contents):
        objs.append(
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources "
            + (fonts % font_obj) + b" /Contents %d 0 R >>" % (4 + i * 2))
        objs.append(b"<< /Length %d >>\nstream\n" % len(stream) + stream
                    + b"\nendstream")
    objs.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    out, offsets = bytearray(b"%PDF-1.4\n"), []
    for i, obj in enumerate(objs, 1):
        offsets.append(len(out))
        out += b"%d 0 obj\n" % i + obj + b"\nendobj\n"

    xref = len(out)
    out += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objs) + 1)
    for off in offsets:
        out += b"%010d 00000 n \n" % off
    out += (b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n"
            % (len(objs) + 1, xref))
    return bytes(out)


def docx_bytes(paragraphs=("Quarterly report", "Revenue grew by 12 percent."),
               table_rows=()) -> bytes:
    doc = DocxDocument()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, cell in enumerate(row):
                table.cell(r, c).text = cell
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def html_bytes(body="<p>Visible text</p>",
               head="<style>.a{color:red}</style>",
               script="<script>var secret = 1;</script>") -> bytes:
    return (f"<!DOCTYPE html><html><head>{head}</head>"
            f"<body>{body}{script}</body></html>").encode()


def es_hit(doc_id="d1", title="Refund Policy", score=8.41,
           snippet="request <em>refunds</em> within 30 days",
           metadata=None):
    hit = {"_score": score,
           "_source": {"doc_id": doc_id, "title": title,
                       "metadata": metadata or {"dept": "finance"}}}
    if snippet is not None:
        hit["highlight"] = {"body": [snippet]}
    return hit


def es_result(hits=(), total=None, aggregations=None):
    hits = list(hits)
    body = {"hits": {"total": total if total is not None
                     else {"value": len(hits)}, "hits": hits}}
    if aggregations:
        body["aggregations"] = aggregations
    return body


def doc_row(**overrides):
    row = {
        "doc_id": "11111111-1111-4111-8111-111111111111",
        "tenant": "acme",
        "title": "Refund Policy",
        "body": "customers may request refunds within 30 days",
        "s3_key": None,
        "status": "PENDING",
        "version": 1,
        "metadata": {},
        "created_at": "2026-01-01T00:00:00Z",
        "failure_reason": None,
    }
    row.update(overrides)
    return row
