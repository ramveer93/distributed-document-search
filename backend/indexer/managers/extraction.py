"""Turn uploaded bytes into plain text.

Runs in the indexer, never in the API: parsing is slow and can hang or crash
on malformed input, so it belongs in a worker you can kill and scale on its
own — not on the request path, and emphatically not inside Elasticsearch
(the ingest-attachment processor would burn data-node CPU that has to answer
searches in under 500 ms).
"""
import io
import tempfile
from html.parser import HTMLParser

import pypdf
from docx import Document as DocxDocument

# a PDF with pages but essentially no extractable text is a scan. pypdf does
# not error on those — it cheerfully returns "" — so without this check we
# would silently index an empty document that can never be found.
MIN_PDF_TEXT = 50


class UnsupportedType(Exception):
    pass


class NeedsOCR(Exception):
    pass


def sniff(raw: bytes) -> str:
    """Magic bytes, never the declared Content-Type — a client mislabelling a
    PDF as text/plain must not put binary in the index."""
    head = raw[:8]
    if head[:4] == b"%PDF":
        return "pdf"
    if head[:4] == b"PK\x03\x04":          # any OOXML container
        return "zip"
    if head[:5] == b"{\\rtf":
        return "rtf"
    stripped = raw[:512].lstrip().lower()
    if stripped.startswith((b"<!doctype html", b"<html")):
        return "html"
    try:
        raw[:4096].decode("utf-8")
        return "text"
    except UnicodeDecodeError:
        return "binary"


def extract(raw: bytes, filename: str = "") -> tuple[str, int]:
    """Returns (text, page_count). Raises UnsupportedType or NeedsOCR."""
    kind = sniff(raw)
    if kind == "pdf":
        return _pdf(raw)
    if kind == "zip":
        if filename.lower().endswith(".docx") or b"word/document.xml" in raw[:8192]:
            return _docx(raw), 0
        return _docx(raw), 0
    if kind == "html":
        return _html(raw.decode("utf-8", errors="replace")), 0
    if kind == "text":
        return raw.decode("utf-8", errors="replace"), 0
    raise UnsupportedType(f"cannot extract text from {kind} content")


def _pdf(raw: bytes) -> tuple[str, int]:
    # streamed to disk rather than held in memory: a PDF's cross-reference
    # table lives at the END of the file, so extraction needs random access to
    # the whole thing — but only one page at a time needs to be resident.
    with tempfile.NamedTemporaryFile(suffix=".pdf") as tmp:
        tmp.write(raw)
        tmp.flush()
        tmp.seek(0)
        try:
            reader = pypdf.PdfReader(tmp)
        except Exception as exc:
            # a corrupt or truncated PDF will never parse, however many times
            # we try. classify it permanent so it skips the retry ladder and
            # lands on FAILED immediately instead of burning three backoffs.
            raise UnsupportedType(f"unreadable pdf: {exc}")
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                raise UnsupportedType("encrypted pdf, no password supplied")
        try:
            pages = [(p.extract_text() or "") for p in reader.pages]
        except Exception as exc:
            raise UnsupportedType(f"unreadable pdf: {exc}")
        text = "\n".join(pages).strip()
        if len(reader.pages) > 0 and len(text) < MIN_PDF_TEXT:
            raise NeedsOCR("no text layer — this document needs OCR")
        return text, len(reader.pages)


def _docx(raw: bytes) -> str:
    doc = DocxDocument(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(t for t in parts if t.strip())


class _Stripper(HTMLParser):
    _SKIP = {"script", "style", "head"}

    def __init__(self):
        super().__init__()
        self.chunks: list[str] = []
        self._skipping = 0

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP:
            self._skipping += 1

    def handle_endtag(self, tag):
        if tag in self._SKIP and self._skipping:
            self._skipping -= 1

    def handle_data(self, data):
        if not self._skipping and data.strip():
            self.chunks.append(data.strip())


def _html(markup: str) -> str:
    s = _Stripper()
    s.feed(markup)
    return "\n".join(s.chunks)
